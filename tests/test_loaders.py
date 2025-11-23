"""Tests for document loaders."""

from pathlib import Path
from typing import List

import pytest

from rag_prep.loaders import (
    CSVLoader,
    HTMLLoader,
    MarkdownLoader,
    TextLoader,
    get_loader,
)
from rag_prep.models import Chunk


def test_text_loader(sample_text_file: Path) -> None:
    """Test the text loader extracts content exactly as written."""
    loader = TextLoader()
    chunks: List[Chunk] = list(loader.load(sample_text_file))

    assert len(chunks) == 1, "Text loader should produce one chunk"
    assert chunks[0].text == "This is a sample text file.\nIt has multiple lines.\nFor testing purposes."
    assert chunks[0].metadata["file_type"] == "text"
    assert chunks[0].metadata["file_name"] == "sample.txt"


def test_markdown_loader(sample_markdown_file: Path) -> None:
    """Test the markdown loader extracts content and metadata."""
    loader = MarkdownLoader()
    chunks: List[Chunk] = list(loader.load(sample_markdown_file))

    assert len(chunks) == 1, "Markdown loader should produce one chunk"
    assert "# Heading" in chunks[0].text, "Should include markdown content"
    assert "This is markdown content" in chunks[0].text
    assert chunks[0].metadata["file_type"] == "markdown"
    assert chunks[0].metadata["file_name"] == "sample.md"


def test_csv_loader(sample_csv_file: Path) -> None:
    """Test the CSV loader treats each row as a document."""
    loader = CSVLoader()
    chunks: List[Chunk] = list(loader.load(sample_csv_file))

    assert len(chunks) == 3, "CSV loader should produce one chunk per row (excluding header)"
    assert "Alice" in chunks[0].text, "First row should contain Alice"
    assert "Bob" in chunks[1].text, "Second row should contain Bob"
    assert "Charlie" in chunks[2].text, "Third row should contain Charlie"

    # Check metadata
    assert chunks[0].metadata["file_type"] == "csv"
    assert chunks[0].metadata["row_index"] == 0
    assert "csv_columns" in chunks[0].metadata


def test_html_loader(sample_html_file: Path) -> None:
    """Test the HTML loader extracts text and removes scripts."""
    loader = HTMLLoader()
    chunks: List[Chunk] = list(loader.load(sample_html_file))

    assert len(chunks) == 1, "HTML loader should produce one chunk"
    assert "Title" in chunks[0].text, "Should extract heading text"
    assert "This is HTML content" in chunks[0].text, "Should extract paragraph text"
    assert "console.log" not in chunks[0].text, "Should remove script content"
    assert chunks[0].metadata["file_type"] == "html"


def test_pdf_loader(temp_dir: Path) -> None:
    """Test the PDF loader if PyPDF2 is available."""
    try:
        import PyPDF2
    except ImportError:
        pytest.skip("PyPDF2 not available")

    # Create a minimal PDF using PyPDF2
    pdf_path = temp_dir / "test.pdf"
    from PyPDF2 import PdfWriter

    writer = PdfWriter()
    page = writer.add_blank_page(width=200, height=200)
    # Note: PyPDF2 doesn't easily add text without a font, so we'll just test the loader exists
    with open(pdf_path, "wb") as f:
        writer.write(f)

    from rag_prep.loaders import PDFLoader

    loader = PDFLoader()
    chunks: List[Chunk] = list(loader.load(pdf_path))

    assert len(chunks) == 1, "PDF loader should produce at least one chunk"
    assert chunks[0].metadata["file_type"] == "pdf"
    assert "num_pages" in chunks[0].metadata


def test_docx_loader(temp_dir: Path) -> None:
    """Test the DOCX loader if python-docx is available."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not available")

    # Create a simple DOCX file
    docx_path = temp_dir / "test.docx"
    doc = Document()
    doc.add_paragraph("This is a test document.")
    doc.add_paragraph("With multiple paragraphs.")
    doc.save(str(docx_path))

    from rag_prep.loaders import DocxLoader

    loader = DocxLoader()
    chunks: List[Chunk] = list(loader.load(docx_path))

    assert len(chunks) == 1, "DOCX loader should produce one chunk"
    assert "This is a test document" in chunks[0].text
    assert "With multiple paragraphs" in chunks[0].text
    assert chunks[0].metadata["file_type"] == "docx"


def test_get_loader_auto_detection(sample_text_file: Path, sample_markdown_file: Path) -> None:
    """Test that get_loader automatically selects the correct loader."""
    text_loader = get_loader(sample_text_file)
    assert isinstance(text_loader, TextLoader), "Should return TextLoader for .txt files"

    md_loader = get_loader(sample_markdown_file)
    assert isinstance(md_loader, MarkdownLoader), "Should return MarkdownLoader for .md files"


def test_get_loader_directory(multi_file_dir: Path) -> None:
    """Test that get_loader returns DirectoryLoader for directories."""
    from rag_prep.loaders import DirectoryLoader

    loader = get_loader(multi_file_dir)
    assert isinstance(loader, DirectoryLoader), "Should return DirectoryLoader for directories"


def test_directory_loader_with_patterns(temp_dir: Path) -> None:
    """Test directory loader respects include/exclude patterns."""
    # Create files with different extensions
    (temp_dir / "file1.txt").write_text("Content 1")
    (temp_dir / "file2.md").write_text("Content 2")
    (temp_dir / "file3.txt").write_text("Content 3")
    (temp_dir / "excluded.txt").write_text("Should not load")

    from rag_prep.loaders import DirectoryLoader

    # Test include pattern
    loader = DirectoryLoader(include_patterns=["file*.txt"])
    chunks = list(loader.load(temp_dir))
    assert len(chunks) == 2, "Should load only files matching include pattern"
    assert all("file1" in c.metadata["source_id"] or "file3" in c.metadata["source_id"] for c in chunks)

    # Test exclude pattern
    loader = DirectoryLoader(exclude_patterns=["excluded.txt"])
    chunks = list(loader.load(temp_dir))
    assert len(chunks) == 3, "Should load all files except excluded"
    assert not any("excluded" in c.metadata["source_id"] for c in chunks)

